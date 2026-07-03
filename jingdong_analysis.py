import json
import csv
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from matplotlib.font_manager import FontProperties
import numpy as np
from datetime import datetime, timedelta
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = "/Users/newcomer/Desktop/workbuddy本地文件"
json_path = f"{BASE}/jingdong_data.json"

with open(json_path, "r", encoding="utf-8") as f:
    raw = json.load(f)

# Sort by trade_date ascending
raw.sort(key=lambda x: x["trade_date"])

import pandas as pd
df = pd.DataFrame(raw)

# Convert trade_date to datetime
df["trade_date_dt"] = pd.to_datetime(df["trade_date"])

# ============================================================
# 1. Save CSV
# ============================================================
csv_data = df[["trade_date", "open", "high", "low", "close", "pre_close", "change", "pct_chg", "vol", "amount"]].copy()
csv_data.columns = ["日期", "开盘价", "最高价", "最低价", "收盘价", "昨收价", "涨跌额", "涨跌幅(%)", "成交量(手)", "成交额(元)"]
csv_path = f"{BASE}/京东方A_每日交易数据.csv"
csv_data.to_csv(csv_path, index=False, encoding="utf-8-sig")
print(f"CSV: {csv_path}")

# ============================================================
# 2. Calculate statistics
# ============================================================
close = df["close"].values
trade_dates = df["trade_date_dt"].values

max_close = df["close"].max()
max_date = df.loc[df["close"].idxmax(), "trade_date_dt"]
min_close = df["close"].min()
min_date = df.loc[df["close"].idxmin(), "trade_date_dt"]
first_close = df.iloc[0]["close"]
last_close = df.iloc[-1]["close"]
year_high = df["high"].max()
year_low = df["low"].min()
avg_close = df["close"].mean()
avg_vol = df["vol"].mean()
total_return = ((last_close - first_close) / first_close) * 100
trading_days = len(df)

# Volatility
daily_returns = df["close"].pct_change().dropna()
annual_volatility = daily_returns.std() * np.sqrt(252) * 100

print(f"Trading days: {trading_days}")
print(f"First close: {first_close:.2f}")
print(f"Last close: {last_close:.2f}")
print(f"Total return: {total_return:.2f}%")

# ============================================================
# 3. Find Chinese font
# ============================================================
font_paths = [
    "/System/Library/Fonts/STSong.ttf",
    "/System/Library/Fonts/PingFang.ttc",
    "/System/Library/Fonts/Supplemental/Songti.ttc",
    "/Library/Fonts/Songti.ttc",
]
chinese_font = None
for fp in font_paths:
    if os.path.exists(fp):
        chinese_font = FontProperties(fname=fp)
        print(f"Using font: {fp}")
        break
if chinese_font is None:
    chinese_font = FontProperties()

# ============================================================
# 4. Plot closing price chart
# ============================================================
fig, ax = plt.subplots(figsize=(14, 6))
ax.plot(df["trade_date_dt"], df["close"], color="#D85A30", linewidth=1.2, label="收盘价")
ax.fill_between(df["trade_date_dt"], df["close"], alpha=0.1, color="#D85A30")

# Annotate max/min
ax.annotate(f"最高: {max_close:.2f}", 
            xy=(max_date, max_close), xytext=(20, 15),
            textcoords="offset points", fontproperties=chinese_font,
            fontsize=10, color="#993C1D",
            arrowprops=dict(arrowstyle="->", color="#993C1D", lw=0.8))
ax.plot(max_date, max_close, "o", color="#993C1D", markersize=5)

ax.annotate(f"最低: {min_close:.2f}",
            xy=(min_date, min_close), xytext=(20, -20),
            textcoords="offset points", fontproperties=chinese_font,
            fontsize=10, color="#0F6E56",
            arrowprops=dict(arrowstyle="->", color="#0F6E56", lw=0.8))
ax.plot(min_date, min_close, "o", color="#0F6E56", markersize=5)

# MA20
df["MA20"] = df["close"].rolling(window=20).mean()
ax.plot(df["trade_date_dt"], df["MA20"], color="#185FA5", linewidth=1, linestyle="--", label="20日均线")

ax.set_title("图1  京东方A(000725.SZ)过去一年每日收盘价走势", fontproperties=chinese_font, fontsize=14, fontweight="bold", pad=15)
ax.set_xlabel("日期", fontproperties=chinese_font, fontsize=11)
ax.set_ylabel("收盘价 (元)", fontproperties=chinese_font, fontsize=11)
ax.xaxis.set_major_locator(mdates.MonthLocator())
ax.xaxis.set_major_formatter(mdates.DateFormatter("%Y-%m"))
plt.setp(ax.get_xticklabels(), rotation=45, ha="right", fontsize=9)
ax.legend(prop=chinese_font, loc="upper left", framealpha=0.9)
ax.grid(True, alpha=0.3, linestyle="--")
ax.spines["top"].set_visible(False)
ax.spines["right"].set_visible(False)
plt.tight_layout()
chart_path = f"{BASE}/京东方A_收盘价走势图.png"
plt.savefig(chart_path, dpi=200, bbox_inches="tight")
plt.close()
print(f"Chart: {chart_path}")

# ============================================================
# 5. Plot volume chart
# ============================================================
fig2, ax2 = plt.subplots(figsize=(14, 3.5))
colors = ["#D85A30" if df.iloc[i]["close"] >= df.iloc[i]["open"] else "#0F6E56" for i in range(len(df))]
ax2.bar(df["trade_date_dt"], df["vol"], color=colors, alpha=0.7, width=1)
ax2.set_title("图2  京东方A(000725.SZ)过去一年每日成交量", fontproperties=chinese_font, fontsize=14, fontweight="bold", pad=15)
ax2.set_xlabel("日期", fontproperties=chinese_font, fontsize=11)
ax2.set_ylabel("成交量 (手)", fontproperties=chinese_font, fontsize=11)
ax2.xaxis.set_major_locator(mdates.MonthLocator())
ax2.xaxis.set_major_formatter(mdates.DateFormatter("%Y-%m"))
plt.setp(ax2.get_xticklabels(), rotation=45, ha="right", fontsize=9)
ax2.grid(True, alpha=0.3, linestyle="--")
ax2.spines["top"].set_visible(False)
ax2.spines["right"].set_visible(False)
plt.tight_layout()
vol_chart_path = f"{BASE}/京东方A_成交量图.png"
plt.savefig(vol_chart_path, dpi=200, bbox_inches="tight")
plt.close()
print(f"Vol Chart: {vol_chart_path}")

# ============================================================
# 6. Generate Word report
# ============================================================
print("\nGenerating Word report...")

doc = Document()

style = doc.styles["Normal"]
font = style.font
font.name = "宋体"
font.size = Pt(10.5)
pf = style.paragraph_format
pf.line_spacing = 1.5
pf.space_before = Pt(0)
pf.space_after = Pt(0)
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_para(text, bold=False, size=10.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "宋体"
    run.font.size = Pt(size)
    run.bold = bold
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return p

def add_heading_custom(text, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "宋体"
    run.font.size = Pt(size)
    run.bold = True
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return p

def add_image_doc(path, width=5.5, caption=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    if caption:
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.line_spacing = 1.5
        cap_p.paragraph_format.space_before = Pt(0)
        cap_p.paragraph_format.space_after = Pt(0)
        cap_run = cap_p.add_run(caption)
        cap_run.font.name = "宋体"
        cap_run.font.size = Pt(10.5)
        cap_run.bold = True
        cap_run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

start_str = df.iloc[0]["trade_date"]
end_str = df.iloc[-1]["trade_date"]
start_fmt = f"{start_str[:4]}年{start_str[4:6]}月{start_str[6:8]}日"
end_fmt = f"{end_str[:4]}年{end_str[4:6]}月{end_str[6:8]}日"

# Title
add_para("京东方A（000725.SZ）股票交易数据分析报告", bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(f"数据区间：{start_fmt} 至 {end_fmt}", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("")

# Section 1
add_heading_custom("一、数据概述", 14)
add_para(
    f"本报告基于京东方A（股票代码：000725.SZ）在深圳证券交易所的每日交易数据，"
    f"数据区间为 {start_fmt} 至 {end_fmt}，共计 {trading_days} 个交易日。"
    f"数据来源为 Tushare Pro，包含每日开盘价、最高价、最低价、收盘价及成交量等字段。"
    f"原始数据已保存为 CSV 格式文件（京东方A_每日交易数据.csv），供后续分析使用。"
)

# Section 2
add_heading_custom("二、关键统计指标", 14)
add_para(
    f"在分析区间内，京东方A股票的关键统计指标如下："
    f"期初收盘价为 ¥{first_close:.2f}，期末收盘价为 ¥{last_close:.2f}，"
    f"区间累计涨跌幅为 {total_return:+.2f}%。"
    f"期间最高价达到 ¥{year_high:.2f}（出现在 {max_date.strftime('%Y年%m月%d日')}），"
    f"最低价下探至 ¥{year_low:.2f}（出现在 {min_date.strftime('%Y年%m月%d日')}），"
    f"日均收盘价为 ¥{avg_close:.2f}，日均成交量约 {avg_vol/1e4:.2f} 万手。"
    f"年化波动率约为 {annual_volatility:.2f}%，反映了该股票在分析区间内的价格波动程度。"
)

# Section 3
add_heading_custom("三、收盘价走势分析", 14)
add_image_doc(chart_path, width=5.5, caption="图1  京东方A(000725.SZ)每日收盘价走势图")

if total_return > 0:
    add_para(
        f"从图1可以看出，京东方A股价在过去一年中整体呈现上涨趋势。"
        f"股价从期初的 ¥{first_close:.2f} 上涨至期末的 ¥{last_close:.2f}，"
        f"累计涨幅为 {total_return:+.2f}%。"
        f"其中，{max_date.strftime('%Y年%m月%d日')}达到区间最高收盘价 ¥{max_close:.2f}，"
        f"而 {min_date.strftime('%Y年%m月%d日')}跌至区间最低收盘价 ¥{min_close:.2f}。"
        f"20日移动平均线（图中蓝色虚线）显示了股价的中期趋势方向。"
    )
else:
    add_para(
        f"从图1可以看出，京东方A股价在过去一年中整体呈现下跌趋势。"
        f"股价从期初的 ¥{first_close:.2f} 下跌至期末的 ¥{last_close:.2f}，"
        f"累计跌幅为 {total_return:+.2f}%。"
        f"其中，{max_date.strftime('%Y年%m月%d日')}达到区间最高收盘价 ¥{max_close:.2f}，"
        f"而 {min_date.strftime('%Y年%m月%d日')}跌至区间最低收盘价 ¥{min_close:.2f}。"
        f"20日移动平均线（图中蓝色虚线）显示了股价的中期趋势方向，"
        f"投资者可结合均线位置判断短期趋势的强弱。"
    )

add_para(
    f"最高价与最低价之间的振幅达到 {((max_close - min_close) / min_close * 100):.2f}%，"
    f"说明期间市场波动较为剧烈。"
)

# Section 4
add_heading_custom("四、成交量分析", 14)
add_image_doc(vol_chart_path, width=5.5, caption="图2  京东方A(000725.SZ)每日成交量图")

add_para(
    f"图2展示了分析区间内京东A每日成交量的变化情况。"
    f"红色柱体表示当日收盘价高于开盘价（上涨日），绿色柱体表示收盘价低于开盘价（下跌日）。"
    f"日均成交量约为 {avg_vol/1e4:.2f} 万手。"
    f"一般而言，成交量放大伴随价格上涨通常被视为多头信号，"
    f"而成交量放大伴随价格下跌则可能预示空头力量增强。"
)

# Section 5
add_heading_custom("五、风险与波动分析", 14)
add_para(
    f"在分析区间内，京东方A股票的年化波动率约为 {annual_volatility:.2f}%。"
    f"波动率是衡量股票价格波动程度的重要指标，年化波动率越高，说明股票价格的不确定性越大，投资风险也相应越高。"
    f"期间最高价与最低价之间的振幅为 {((year_high - year_low) / year_low * 100):.2f}%。"
)

# Section 6
add_heading_custom("六、总结", 14)
add_para(
    f"综上所述，京东方A（000725.SZ）在过去一年的交易中，"
    f"股价从 ¥{first_close:.2f} 变动至 ¥{last_close:.2f}，"
    f"累计涨跌幅 {total_return:+.2f}%，年化波动率 {annual_volatility:.2f}%。"
    f"作为国内面板行业的龙头企业，京东方A的股价走势受到面板行业周期、"
    f"OLED 业务进展、市场需求变化以及宏观经济环境等多重因素影响。"
    f"本报告仅基于历史交易数据进行统计分析，不构成任何投资建议。"
)

add_para("")
add_para("免责声明：本报告中的数据和分析仅供学习参考，不构成任何投资建议。投资有风险，入市需谨慎。")

docx_path = f"{BASE}/京东方A_股票分析报告.docx"
doc.save(docx_path)
print(f"Word: {docx_path}")

# ============================================================
# 7. Generate HTML dashboard (ai-quant style)
# ============================================================
print("\nGenerating HTML dashboard...")

# Data prep
data = raw  # already sorted
total_days = len(data)
last = data[-1]

week_high = max(r["high"] for r in data)
week_low = min(r["low"] for r in data)
week_high_item = max(data, key=lambda r: r["high"])
week_low_item = min(data, key=lambda r: r["low"])
week_high_date = week_high_item["trade_date"]
week_low_date = week_low_item["trade_date"]
daily_change = last["pct_chg"]
avg_vol_all = sum(r["vol"] for r in data) / total_days

# Recent 10
recent_10 = data[-10:][::-1]

dates_json = json.dumps([r["trade_date"] for r in data])
opens_json = json.dumps([r["open"] for r in data])
highs_json = json.dumps([r["high"] for r in data])
lows_json = json.dumps([r["low"] for r in data])
closes_json = json.dumps([r["close"] for r in data])
volumes_json = json.dumps([r["vol"] for r in data])
normalized_json = json.dumps([r["close"] / data[0]["close"] * 100 for r in data])

recent_rows = ""
for r in recent_10:
    d = r["trade_date"]
    d_fmt = f"{d[:4]}-{d[4:6]}-{d[6:8]}"
    cls = "up" if r["pct_chg"] >= 0 else "down"
    sign = "+" if r["pct_chg"] >= 0 else ""
    recent_rows += f"""<tr>
      <td>{d_fmt}</td>
      <td>{r["open"]:.2f}</td>
      <td>{r["high"]:.2f}</td>
      <td>{r["low"]:.2f}</td>
      <td class="{cls}">{r["close"]:.2f}</td>
      <td class="{cls}">{sign}{r["pct_chg"]:.2f}%</td>
      <td>{(r["vol"]/1e4):.2f}万</td>
    </tr>"""

html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>京东方A(000725.SZ) 股票分析看板</title>
<script src="https://cdn.jsdelivr.net/npm/echarts@5.5.0/dist/echarts.min.js"></script>
<style>
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'PingFang SC', 'Microsoft YaHei', sans-serif; background: #0d1117; color: #e6edf3; }}
.header {{ background: linear-gradient(135deg, #161b22 0%, #0d1117 100%); border-bottom: 1px solid #30363d; padding: 32px 40px 28px; }}
.header h1 {{ font-size: 28px; font-weight: 700; }}
.header .subtitle {{ color: #8b949e; font-size: 14px; margin-top: 6px; }}
.header .badge {{ display: inline-block; background: #1f6feb22; color: #58a6ff; border: 1px solid #1f6feb44; border-radius: 20px; padding: 2px 12px; font-size: 12px; margin-top: 8px; }}
.container {{ max-width: 1200px; margin: 0 auto; padding: 24px 20px; }}
.stats-row {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(160px, 1fr)); gap: 16px; margin-bottom: 28px; }}
.stat-card {{ background: #161b22; border: 1px solid #30363d; border-radius: 12px; padding: 18px 20px; }}
.stat-card .label {{ font-size: 12px; color: #8b949e; margin-bottom: 6px; text-transform: uppercase; letter-spacing: 0.5px; }}
.stat-card .value {{ font-size: 24px; font-weight: 700; }}
.stat-card .value.up {{ color: #da3633; }}
.stat-card .value.down {{ color: #238636; }}
.stat-card .sub {{ font-size: 12px; color: #8b949e; margin-top: 4px; }}
.section {{ margin-bottom: 28px; }}
.section-title {{ font-size: 18px; font-weight: 600; margin-bottom: 14px; padding-bottom: 8px; border-bottom: 1px solid #30363d; }}
.chart-box {{ background: #161b22; border: 1px solid #30363d; border-radius: 12px; padding: 20px; margin-bottom: 20px; }}
.chart-box .chart-title {{ font-size: 14px; font-weight: 600; color: #e6edf3; margin-bottom: 12px; }}
#klineChart {{ height: 480px; }}
#volumeChart {{ height: 160px; }}
.metrics-table {{ width: 100%; border-collapse: collapse; font-size: 14px; }}
.metrics-table th {{ text-align: left; padding: 10px 14px; background: #161b22; border-bottom: 1px solid #30363d; color: #8b949e; font-weight: 500; }}
.metrics-table td {{ padding: 10px 14px; border-bottom: 1px solid #21262d; }}
.metrics-table tr:hover td {{ background: #1c2128; }}
.metrics-table .hl {{ font-weight: 600; }}
.data-table {{ width: 100%; border-collapse: collapse; font-size: 13px; }}
.data-table th {{ text-align: left; padding: 10px 14px; background: #161b22; border-bottom: 1px solid #30363d; color: #8b949e; font-weight: 500; font-size: 12px; }}
.data-table td {{ padding: 10px 14px; border-bottom: 1px solid #21262d; }}
.data-table tr:hover td {{ background: #1c2128; }}
.data-table .up {{ color: #da3633; }}
.data-table .down {{ color: #238636; }}
.note {{ background: #161b22; border: 1px solid #30363d; border-radius: 12px; padding: 16px 20px; font-size: 13px; color: #8b949e; line-height: 1.7; }}
.note strong {{ color: #e6edf3; }}
.footer {{ text-align: center; color: #484f58; font-size: 12px; padding: 24px; border-top: 1px solid #21262d; }}
</style>
</head>
<body>

<div class="header">
  <h1>📈 京东方A (000725.SZ)</h1>
  <div class="subtitle">深圳证券交易所主板 · 过去一年行情分析 · {data[0]["trade_date"][:4]}-{data[0]["trade_date"][4:6]}-{data[0]["trade_date"][6:8]} 至 {data[-1]["trade_date"][:4]}-{data[-1]["trade_date"][4:6]}-{data[-1]["trade_date"][6:8]} · 共 {total_days} 个交易日</div>
  <div class="badge">数据来源: Tushare Pro</div>
</div>

<div class="container">

  <div class="stats-row">
    <div class="stat-card">
      <div class="label">最新收盘</div>
      <div class="value">{last["close"]:.2f}</div>
      <div class="sub">CNY</div>
    </div>
    <div class="stat-card">
      <div class="label">日涨跌幅</div>
      <div class="value {'up' if daily_change >= 0 else 'down'}">{'+' if daily_change >= 0 else ''}{daily_change:.2f}%</div>
      <div class="sub">{'涨' if daily_change >= 0 else '跌'}</div>
    </div>
    <div class="stat-card">
      <div class="label">区间涨跌幅</div>
      <div class="value {'up' if total_return >= 0 else 'down'}">{'+' if total_return >= 0 else ''}{total_return:.2f}%</div>
      <div class="sub">过去一年</div>
    </div>
    <div class="stat-card">
      <div class="label">区间最高</div>
      <div class="value">{week_high:.2f}</div>
      <div class="sub">{week_high_date[:4]}-{week_high_date[4:6]}-{week_high_date[6:8]}</div>
    </div>
    <div class="stat-card">
      <div class="label">区间最低</div>
      <div class="value">{week_low:.2f}</div>
      <div class="sub">{week_low_date[:4]}-{week_low_date[4:6]}-{week_low_date[6:8]}</div>
    </div>
    <div class="stat-card">
      <div class="label">日均成交量</div>
      <div class="value">{(avg_vol_all/1e4):.0f}万</div>
      <div class="sub">手</div>
    </div>
  </div>

  <div class="section">
    <div class="section-title">📊 K 线走势</div>
    <div class="chart-box">
      <div class="chart-title">京东方A 日 K 线图</div>
      <div id="klineChart"></div>
    </div>
  </div>

  <div class="section">
    <div class="section-title">📊 成交量</div>
    <div class="chart-box">
      <div class="chart-title">每日成交量</div>
      <div id="volumeChart"></div>
    </div>
  </div>

  <div class="section">
    <div class="section-title">📊 价格走势对比</div>
    <div class="chart-box">
      <div class="chart-title">收盘价走势 (归一化至 100)</div>
      <div id="trendChart"></div>
    </div>
  </div>

  <div class="section">
    <div class="section-title">📋 关键指标汇总</div>
    <div class="chart-box">
      <table class="metrics-table">
        <tr><th>指标</th><th>京东方A (000725.SZ)</th></tr>
        <tr><td>最新收盘价</td><td class="hl">{last["close"]:.2f} CNY</td></tr>
        <tr><td>开盘价</td><td>{last["open"]:.2f} CNY</td></tr>
        <tr><td>日内最高</td><td>{last["high"]:.2f} CNY</td></tr>
        <tr><td>日内最低</td><td>{last["low"]:.2f} CNY</td></tr>
        <tr><td>昨收价</td><td>{last["pre_close"]:.2f} CNY</td></tr>
        <tr><td>涨跌额</td><td class="{'up' if last['change'] >= 0 else 'down'}">{'+' if last['change'] >= 0 else ''}{last["change"]:.2f} CNY</td></tr>
        <tr><td>涨跌幅</td><td class="{'up' if last['pct_chg'] >= 0 else 'down'}">{'+' if last['pct_chg'] >= 0 else ''}{last["pct_chg"]:.2f}%</td></tr>
        <tr><td>成交量</td><td>{last["vol"]:,.0f} 手</td></tr>
        <tr><td>成交额</td><td>{(last["amount"]/1e8):.2f}亿 CNY</td></tr>
        <tr><td>过去一年涨跌幅</td><td class="{'up' if total_return >= 0 else 'down'}">{'+' if total_return >= 0 else ''}{total_return:.2f}%</td></tr>
        <tr><td>区间最高</td><td>{week_high:.2f} CNY</td></tr>
        <tr><td>区间最低</td><td>{week_low:.2f} CNY</td></tr>
        <tr><td>区间振幅</td><td>{((max_close - min_close) / min_close * 100):.2f}%</td></tr>
        <tr><td>总交易日</td><td>{total_days} 天</td></tr>
      </table>
    </div>
  </div>

  <div class="section">
    <div class="section-title">📋 最近 10 个交易日</div>
    <div class="chart-box">
      <table class="data-table">
        <tr><th>日期</th><th>开盘</th><th>最高</th><th>最低</th><th>收盘</th><th>涨跌幅</th><th>成交量</th></tr>
        {recent_rows}
      </table>
    </div>
  </div>

  <div class="section">
    <div class="section-title">⚠️ 注意事项</div>
    <div class="note">
      <strong>数据说明：</strong>本看板数据来源于 Tushare Pro A股日线接口，更新频率为每日收盘后更新。<br>
      <strong>风险提示：</strong>所有数据仅供参考，不构成任何投资建议。投资有风险，入市需谨慎。<br>
      <strong>配色说明：</strong>A股通用配色——<span style="color:#da3633">红色代表上涨</span>，<span style="color:#238636">绿色代表下跌</span>。
    </div>
  </div>

</div>

<div class="footer">Made with WorkBuddy · 数据来源 Tushare Pro</div>

<script>
var dates = {dates_json};
var klineData = [];
for (var i = 0; i < dates.length; i++) {{
  klineData.push([{opens_json}[i], {closes_json}[i], {lows_json}[i], {highs_json}[i]]);
}}
var volColors = [];
for (var i = 0; i < dates.length; i++) {{
  volColors.push({closes_json}[i] >= {opens_json}[i] ? '#da3633' : '#238636');
}}

// K-line
var klineChart = echarts.init(document.getElementById('klineChart'));
klineChart.setOption({{
  tooltip: {{
    trigger: 'axis', axisPointer: {{ type: 'cross' }},
    backgroundColor: '#1c2128', borderColor: '#30363d', textStyle: {{ color: '#e6edf3' }},
    formatter: function(params) {{
      var p = params[0], d = p.axisValue, idx = dates.indexOf(d);
      if (idx === -1) return '';
      return '<b>' + d.slice(0,4)+'-'+d.slice(4,6)+'-'+d.slice(6,8)+'</b><br/>开盘: '+{opens_json}[idx].toFixed(2)+'<br/>收盘: '+{closes_json}[idx].toFixed(2)+'<br/>最高: '+{highs_json}[idx].toFixed(2)+'<br/>最低: '+{lows_json}[idx].toFixed(2);
    }}
  }},
  grid: {{ left: '4%', right: '4%', top: '6%', bottom: '6%' }},
  xAxis: {{ type: 'category', data: dates, axisLabel: {{ rotate: 45, fontSize: 11, color: '#8b949e', interval: 'auto' }}, axisLine: {{ lineStyle: {{ color: '#30363d' }} }}, splitLine: {{ show: false }} }},
  yAxis: {{ type: 'value', scale: true, splitLine: {{ lineStyle: {{ type: 'dashed', color: '#21262d' }} }}, axisLabel: {{ color: '#8b949e' }} }},
  dataZoom: [ {{ type: 'inside', start: 50, end: 100 }}, {{ type: 'slider', start: 50, end: 100, height: 20, bottom: 0, borderColor: '#30363d', backgroundColor: '#161b22', fillerColor: '#1f6feb33' }} ],
  series: [{{ type: 'candlestick', data: klineData, itemStyle: {{ color: '#da3633', color0: '#238636', borderColor: '#da3633', borderColor0: '#238636' }} }}]
}});

// Volume
var volChart = echarts.init(document.getElementById('volumeChart'));
volChart.setOption({{
  tooltip: {{ trigger: 'axis', backgroundColor: '#1c2128', borderColor: '#30363d', textStyle: {{ color: '#e6edf3' }}, formatter: function(p) {{ return p[0].axisValue.slice(0,4)+'-'+p[0].axisValue.slice(4,6)+'-'+p[0].axisValue.slice(6,8)+'<br/>成交量: '+(p[0].value/1e4).toFixed(0)+'万手'; }} }},
  grid: {{ left: '4%', right: '4%', top: '6%', bottom: '12%' }},
  xAxis: {{ type: 'category', data: dates, axisLabel: {{ show: false }}, axisLine: {{ lineStyle: {{ color: '#30363d' }} }}, splitLine: {{ show: false }} }},
  yAxis: {{ type: 'value', splitLine: {{ lineStyle: {{ type: 'dashed', color: '#21262d' }} }}, axisLabel: {{ color: '#8b949e', formatter: function(v) {{ return (v/1e4).toFixed(0)+'万'; }} }} }},
  dataZoom: [ {{ type: 'inside', start: 50, end: 100 }}, {{ type: 'slider', start: 50, end: 100, height: 16, bottom: 0, borderColor: '#30363d', backgroundColor: '#161b22', fillerColor: '#1f6feb33' }} ],
  series: [{{ type: 'bar', data: volColors.map(function(c,i) {{ return {{ value: {volumes_json}[i], itemStyle: {{ color: c }} }}; }}) }}]
}});

// Trend
var trendChart = echarts.init(document.getElementById('trendChart'));
trendChart.setOption({{
  tooltip: {{ trigger: 'axis', backgroundColor: '#1c2128', borderColor: '#30363d', textStyle: {{ color: '#e6edf3' }}, formatter: function(p) {{ return '<b>'+p[0].axisValue.slice(0,4)+'-'+p[0].axisValue.slice(4,6)+'-'+p[0].axisValue.slice(6,8)+'</b><br/>归一化: '+p[0].value.toFixed(2); }} }},
  grid: {{ left: '4%', right: '4%', top: '6%', bottom: '6%' }},
  xAxis: {{ type: 'category', data: dates, axisLabel: {{ rotate: 45, fontSize: 11, color: '#8b949e', interval: 'auto' }}, axisLine: {{ lineStyle: {{ color: '#30363d' }} }}, splitLine: {{ show: false }} }},
  yAxis: {{ type: 'value', splitLine: {{ lineStyle: {{ type: 'dashed', color: '#21262d' }} }}, axisLabel: {{ color: '#8b949e' }} }},
  dataZoom: [ {{ type: 'inside', start: 50, end: 100 }}, {{ type: 'slider', start: 50, end: 100, height: 16, bottom: 0, borderColor: '#30363d', backgroundColor: '#161b22', fillerColor: '#1f6feb33' }} ],
  series: [{{ type: 'line', data: {normalized_json}, smooth: true, lineStyle: {{ color: '#58a6ff', width: 2 }}, areaStyle: {{ color: {{ type: 'linear', x: 0, y: 0, x2: 0, y2: 1, colorStops: [{{ offset: 0, color: '#58a6ff66' }}, {{ offset: 1, color: '#58a6ff00' }}] }} }}, symbol: 'none' }}]
}});

window.addEventListener('resize', function() {{ klineChart.resize(); volChart.resize(); trendChart.resize(); }});
</script>
</body>
</html>"""

html_path = f"{BASE}/京东方A_K线看板.html"
with open(html_path, "w", encoding="utf-8") as f:
    f.write(html)
print(f"HTML: {html_path}")

print("\n=== ALL DONE ===")
print(f"1. CSV:  {csv_path}")
print(f"2. Chart: {chart_path}")
print(f"3. Vol Chart: {vol_chart_path}")
print(f"4. Word: {docx_path}")
print(f"5. HTML: {html_path}")
